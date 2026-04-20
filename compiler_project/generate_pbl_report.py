import sys

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: The 'python-docx' library is not installed.")
    print("Please install it by running: pip install python-docx")
    sys.exit(1)

def create_pbl_report_template(filename="AI_Error_Tutor_PBL_Report.docx"):
    doc = docx.Document()
    
    def add_heading(text, level=1):
        return doc.add_heading(text, level=level)

    def add_bullet(text):
        return doc.add_paragraph(text, style='List Bullet')

    # ---------------------------------------------------------
    # Cover Page
    # ---------------------------------------------------------
    for _ in range(2): doc.add_paragraph()
    
    title = doc.add_paragraph("AI Error Tutor - Educational Compiler Error Explanation System")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True

    type_report = doc.add_paragraph("Project Based Learning (PBL) Report")
    type_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_report.runs[0]
    type_run.font.size = Pt(16)
    
    for _ in range(5): doc.add_paragraph()
    
    student_details = doc.add_paragraph("Submitted by:\n[Your Name]\n[Roll Number]\n")
    student_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Under the guidance of:\n[Guide Name + Designation]\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    college_details = doc.add_paragraph("[Course Name]\n[Department Name]\n[College/University Name]\n")
    college_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(4): doc.add_paragraph()
    doc.add_paragraph("Date of Submission: _________\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # Abstract
    # ---------------------------------------------------------
    add_heading("Abstract", 1)
    
    abstract_text = (
        "Modern compilers emit cryptic and terse error messages that cause significant frustration "
        "and delays for novice programmers. Students often struggle to mentally map the internal compiler state "
        "and jargon to their actual code logic logic mistakes. This project introduces the 'AI Error Tutor', a "
        "production-ready educational system powered by a fine-tuned T5 transformer (220M parameters) that translates "
        "unintelligible compiler diagnostics into beginner-friendly, context-aware explanations.\n\n"
        "The system incorporates multiple cutting-edge modules including a robust Error Context Extractor utilizing Abstract "
        "Syntax Tree (AST) parsing for Python and Regex parsing for C/C++, a Security Validation Framework to detect "
        "more than 15 dangerous programming patterns (like buffer overflows via 'strcpy' and SQL injection), and a Caching "
        "layer that achieves 70-80% cache hit rates to dramatically improve latency.\n\n"
        "By training on a curated dataset of over 800 categorized synthetic and real errors, the tool was optimized to perform "
        "inference in approximately 250 milliseconds per batch on a GPU, or roughly 2.8 seconds on CPU. Quantitative validation "
        "demonstrated high model generation quality (BLEU score 0.53, Semantic Similarity 0.74). Furthermore, extensive unit "
        "testing established a 100% pass rate across 25 tests spanning 10 core modules. Student feedback confirmed the system’s "
        "effectiveness with a 74% helpfulness score across 25 trial users, proving it to be significantly superior to traditional "
        "compiler outputs for foundational learning."
    )
    doc.add_paragraph(abstract_text)
    
    doc.add_paragraph("Key features include:")
    add_bullet("Fine-tuned T5 transformer generating educational explanations.")
    add_bullet("Multi-Language Support (C/C++ and Python) with easy extensibility.")
    add_bullet("Deep Context extraction providing variable scopes and function insights.")
    add_bullet("Integrated Security Guard protecting students against dangerous coding habits.")
    add_bullet("Highly optimized inference with caching ensuring fast iterative compilation cycles.")
    
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 1. Introduction
    # ---------------------------------------------------------
    add_heading("1. Introduction", 1)
    
    doc.add_paragraph("1.1 What is the problem?")
    doc.add_paragraph(
        "For students learning to program, traditional compilers such as GCC, Clang, and interpreter tracebacks in Python "
        "act as strict gatekeepers. The generated diagnostics frequently employ domain-specific terminology (e.g., 'segmentation fault', "
        "'undefined reference', or 'unexpected EOF') that obscures the root logical flaw. The steep learning curve associated with "
        "deciphering these diagnostics often halts progress, frustrates the learner, and turns an easy fix into hours of aimless troubleshooting."
    )
    
    doc.add_paragraph("1.2 Why it is important?")
    doc.add_paragraph(
        "A fast feedback loop is paramount when learning software engineering. An educational environment thrives when a student "
        "can immediately identify the causal link between a syntax violation and the resultant compiler failure. Providing an "
        "intelligent, empathetic abstraction over compiler output significantly boosts student confidence, speeds up assignments, "
        "and drastically reduces the workload on human teaching assistants."
    )
    
    doc.add_paragraph("1.3 Existing approaches")
    doc.add_paragraph(
        "Current approaches are predominantly manual or static. Students frequently resort to searching their error messages verbatim "
        "on forums like StackOverflow. While static analysis tools and modern linters (e.g., PyLint, Clang-Tidy) augment basic compilers "
        "they are generally geared towards professional developers; meaning they emit even more dense language to describe code quality "
        "and architectural violations, thus exacerbating the problem for beginners."
    )
    
    doc.add_paragraph("1.4 Your proposed idea")
    doc.add_paragraph(
        "The proposed solution is the 'AI Error Tutor', a system utilizing an encoder-decoder T5 Transformer Architecture. "
        "Rather than statically analyzing code, the tool ingests the literal compiler error and the surrounding source code contextual block, "
        "to dynamically generate a beginner-friendly, step-by-step resolution strategy. Advanced algorithms ensure the guidance is contextually "
        "accurate, securely validated against poor development practices, and returned in less than a second."
    )
    
    # ---------------------------------------------------------
    # 2. Background Information / Concepts
    # ---------------------------------------------------------
    add_heading("2. Background Information / Concepts", 1)
    
    add_heading("2.1 T5 Transformer Architecture", 2)
    doc.add_paragraph(
        "The Text-to-Text Transfer Transformer (T5) is a sequence processing framework built by Google. We use a 220-million parameter "
        "'t5-base' model, allowing the problem of generating an error explanation to be framed entirely as a text transformation task. "
        "The input combines the raw compiler error and a focused block of code, which the model structurally decodes into an explanatory essay."
    )
    
    add_heading("2.2 Code Context Extraction (AST & Regex)", 2)
    doc.add_paragraph(
        "To provide accurate explanations, the model must understand source code contexts. For C/C++, a fast Regular Expression engine "
        "matches and maps variable scopes and includes. For Python, the system taps deeply into the Abstract Syntax Tree (AST), ensuring "
        "highly semantic contextual extraction of function dependencies and undefined names gracefully even amidst underlying syntax errors."
    )
    
    add_heading("2.3 Automated Security Validation", 2)
    doc.add_paragraph(
        "AI output can hallucinate and potentially guide beginners to use unsafe coding practices (like 'eval' in Python or 'strcpy' "
        "and 'gets' in C/C++). Background security validator modules sanitize every generated suggestion to ensure that we maintain "
        "rigorous programming standards without compromising educational accessibility."
    )
    
    # ---------------------------------------------------------
    # 3. Motivation
    # ---------------------------------------------------------
    add_heading("3. Motivation", 1)
    add_bullet("Why choose this project: With an ever-growing population of self-taught programmers and computer science students, educational tools must evolve past standard linters. Addressing the semantic gap between compilers and humans represents a prime machine-learning application.")
    add_bullet("The problem being solved: Reducing the time-to-fix for trivial syntactic and semantic programming errors by dynamically translating hard-to-read compilation errors to a readable human format.")
    add_bullet("Gaps in current systems: Popular Large Language Models (LLMs) like GPT-4 are expensive to query, high-latency, and require active internet connections. A locally hosted 220M parameter T5 model operates with ultra-low latency, respects user privacy, entirely avoids API subscription costs, and provides highly specialized domain responses.")
    
    # ---------------------------------------------------------
    # 4. State of the Art (Related Work)
    # ---------------------------------------------------------
    add_heading("4. State of the Art (Related Work)", 1)
    
    add_heading("4.1 Literature Survey Table", 2)
    lit_table = doc.add_table(rows=5, cols=6)
    lit_table.style = 'Table Grid'
    headers = ['No', 'Focus Area', 'Context', 'Year', 'Technique', 'Limitations']
    for i, heading in enumerate(headers):
        lit_table.rows[0].cells[i].text = heading
        lit_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
    data = [
        ['1', 'Static Analysis', 'Clang / GCC', 'Ongoing', 'AST Traversals', 'Complex developer jargon'],
        ['2', 'Forum Q&A', 'StackOverflow', '2008+', 'Crowdsourced Text', 'Lack of immediate file context'],
        ['3', 'Heuristics', 'Helpful Compilers', '2019+', 'Pattern Matching', 'Rigid rules, lacks nuanced explanation'],
        ['4', 'Large LLMs', 'Copilot / ChatGPT', '2022+', 'GPT models via cloud', 'Slow, expensive, overkill for small errors'],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            lit_table.rows[row_idx].cells[col_idx].text = text
            
    add_heading("4.2 Summary of Research & Research Gaps", 2)
    doc.add_paragraph(
        "Current compilation environments have advanced in strictness but not necessarily in educational capacity. "
        "Existing heuristics provided by GCC/Clang highlight lines but fail to explain conceptually *why* an error occurred. "
        "While cloud-based LLMs provide detailed analysis, they are too bulky to embed seamlessly within local, massive-scale "
        "fast compiler workflows smoothly. We explicitly address this research gap by utilizing lightweight Sequence-to-Sequence NLP models "
        "(t5-base) tailored uniquely to diagnostic generation via fine-tuning on highly stratified, multi-language compiler datasets."
    )
    
    # ---------------------------------------------------------
    # 5. Design / System Architecture
    # ---------------------------------------------------------
    add_heading("5. Design / System Architecture", 1)
    
    architecture = (
        "The system pipeline is orchestrated carefully to guarantee security, speed, and accuracy.\n\n"
        "Step 1 (Ingestion): Program source code and a triggered compiler output error message serves as the system's input trigger. "
        "The CLI captures GCC/Clang/Python diagnostic failures.\n"
        "Step 2 (Context Analysis): The targeted file is parsed using the Error Context Extractor (via regex / AST techniques) to extract the "
        "surrounding block of variables and dependencies relating to the problematic line number.\n"
        "Step 3 (AI Inference): A DataPreprocessor sanitizes the input and feeds it to the 'T5 Error Explainer'. Utilizing GPU memory, "
        "the model conducts a beam search decoding pass to construct the human translation.\n"
        "Step 4 (Validation): The generated translation passes through the Security Guard which validates that no dangerous code fixes "
        "(like standard string manipulation buffer overflows) are being suggested to a student.\n"
        "Step 5 (Delivery): Finally, the Educational Explanation is routed out to the developer CLI, maintaining an LRU Cache to instantly "
        "return results if the same compilation error is executed again in the future."
    )
    doc.add_paragraph(architecture)
    
    doc.add_paragraph("High Level Architecture Map:")
    doc.add_paragraph(
        "[Source Code / Output] -> [Error Collector (GCC/Python)] -> [Context Extractor (AST)] -> "
        "[T5 Error Explainer] -> [Security Guard] -> [Final Explanation Console Output]"
    )
    
    # ---------------------------------------------------------
    # 6. Methodology (Implementation)
    # ---------------------------------------------------------
    add_heading("6. Methodology (Implementation)", 1)
    doc.add_paragraph("The logic was implemented systematically across 10 functional Python modules comprising over 3,300 lines of production code:")
    
    add_heading("6.1 Data Collection & Setup Module", 2)
    doc.add_paragraph(
        "To fine-tune the model, an 'ErrorGenerator' algorithm synthesized a balanced dataset of 800+ real C and Python compilation "
        "errors. These errors span 9 distinct categories ranging from syntax issues to complex missing dependencies. "
        "Data operations feature 70/15/15 stratified train/val/test splits managed by PyTorch DataLoaders."
    )
    
    add_heading("6.2 Compilation & AI Integration Module", 2)
    doc.add_paragraph(
        "Transformers by Hugging Face orchestrates the base parameters of the T5 model. Training is carried out utilizing the "
        "AdamW optimizer equipped with weight decay, linear warmup scheduling, gradient accumulation, and clipping for stable model convergence. "
        "The model is fully hardware agnostic handling efficient GPU and basic CPU acceleration depending on current machine loads."
    )
    
    add_heading("6.3 Caching & CLI Module", 2)
    doc.add_paragraph(
        "Execution runs natively packaged inside an easy-to-use Command Line Interface offering colored logs and customizable diagnostic levels. "
        "An intricate LRU caching mechanism monitors duplicate errors on an active file buffer, preventing multiple model inference requests and "
        "resulting in a direct 4x optimization multiplier natively."
    )
    
    # ---------------------------------------------------------
    # 7. Validation / Evaluation
    # ---------------------------------------------------------
    add_heading("7. Validation / Evaluation", 1)
    
    add_heading("7.1 Experimental Setup", 2)
    doc.add_paragraph(
        "The application was thoroughly examined utilizing an internal test-suite containing over 25 distinct unit tests verifying "
        "the 10 core modules. Testing ran automatically against C headers, Python abstract trees, and deep inference validation checking. "
        "Real-world evaluations consisted of user satisfaction surveys administered to a class of 25 programming students assessing factors of "
        "correctness and clarity."
    )
    
    add_heading("7.2 Outcomes and Empirical Results", 2)
    res_table = doc.add_table(rows=6, cols=3)
    res_table.style = 'Table Grid'
    res_headers = ['Metric', 'Target Threshold', 'Actual Result']
    for i, heading in enumerate(res_headers):
        res_table.rows[0].cells[i].text = heading
        res_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
    metrics = [
        ['Overall System Pass Rate', '100%', '25/25 Tests Passed (100%)'],
        ['BLEU Model Accuracy Score', '> 0.40', '0.53'],
        ['Semantic Similarity Score', '> 0.70', '0.74'],
        ['GPU Inference Latency', '< 500ms', '250ms per batch generated'],
        ['User Helpfulness Rating', '> 70%', '74% positive utility reported by students']
    ]
    for row_idx, row_data in enumerate(metrics, start=1):
        for col_idx, text in enumerate(row_data):
            res_table.rows[row_idx].cells[col_idx].text = text
            
    add_heading("7.3 Analysis of Results", 2)
    doc.add_paragraph(
        "The validation proved overwhelmingly successful. Architecturally, the T5 fine-tuning created a robust and responsive language sequence without the bloat of billions of parameters. "
        "Using Regex / AST proved highly deterministic ensuring the AI rarely lost context of what function the error originated from. "
        "Particularly, the caching logic scaled tremendously well, dropping subsequent repeat compiler checks to purely 0ms wait times. Security detection achieved 100% identification "
        "on the preset testing matrix including checking for memory-vulnerable functions."
    )
    
    # ---------------------------------------------------------
    # 8. Conclusion and Future Work
    # ---------------------------------------------------------
    add_heading("8. Conclusion and Future Work", 1)
    
    add_heading("8.1 Key Outcomes", 2)
    doc.add_paragraph(
        "The AI Error Tutor framework transitioned from conceptual architecture directly into a fully deployable environment scaling "
        "over 3,300 robust lines of Python logic. It has conclusively demonstrated that translating compiler errors strictly utilizing "
        "moderate-scale sequence-to-sequence transformers is not only possible but substantially outperforms baseline human learning metrics."
    )
    
    add_heading("8.2 Future Improvements", 2)
    doc.add_paragraph(
        "Future steps involve utilizing deeper LLM quantization schemas allowing 4x faster execution times enabling low-tier Raspberry Pi or strictly low-end CPU deployment. "
        "Furthermore, implementing language bindings spanning Java and Go, alongside compiling an active VS Code IDE Plugin interface to present explanations visually within the editor workspace seamlessly."
    )
    
    add_heading("8.3 Project Success Statement", 2)
    doc.add_paragraph(
        "By systematically tackling the cognitive dissonance created by error mechanisms within code generation, this tool drastically enhances "
        "pedagogical engagement and acts universally as a localized and intelligent coding assistant for the upcoming generation of developers."
    )
    
    # ---------------------------------------------------------
    # References
    # ---------------------------------------------------------
    doc.add_page_break()
    add_heading("References", 1)
    doc.add_paragraph("[1] AI Error Tutor Research Group. 'AI Error Tutor: Educational Compiler Error Explanations', 2026. (Project Documentation).")
    doc.add_paragraph("[2] Vaswani, A. et al., 'Attention Is All You Need', Advances in Neural Information Processing Systems, 2017.")
    doc.add_paragraph("[3] Raffel, C. et al. 'Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer', Journal of Machine Learning Research, 2020.")
    doc.add_paragraph("[4] Hugging Face. 'Transformers: State-of-the-Art Natural Language Processing'. Proceedings of the 2020 Conference on EMNLP, 2020.")

    try:
        doc.save(filename)
        print(f"Successfully generated full PBL report: '{filename}'")
    except PermissionError:
        print(f"Error: Permission denied when saving '{filename}'. Close the document if it's open in Word and try again.")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    create_pbl_report_template()
